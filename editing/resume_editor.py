"""
editing/resume_editor.py
Layer 4 - Editing: Intelligently merge transformed content into the existing resume.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Optional, List, Dict

from docx import Document
from docx.shared import Pt

from core.llm import LLMClient
from core.logger import LayerLogger
from core.models import TransformedProject, EditDecision

_log = LayerLogger("editing")

STYLE_CHECK_SYSTEM = """You are a resume editor checking style consistency.
Given existing resume bullets and a new bullet, rewrite the new bullet to match
the existing style (tense, format, length, tone) if needed.
Return JSON: {"bullet": "...", "changed": true/false, "reason": "..."}"""

# Bullet-like paragraph styles used in DOCX resumes
_BULLET_STYLE_KEYWORDS = {"list", "bullet"}

# Characters that indicate a bullet point prefix
_BULLET_CHARS = set("•·-–—*►▶▸◆◉○▪▫")


def edit(
    projects: List[TransformedProject],
    resume_path: str,
    output_path: str,
    llm: LLMClient,
    min_confidence: float = 0.6,
) -> List[EditDecision]:
    _log.info(f"Opening resume: {resume_path}")
    doc = Document(resume_path)

    structure = _parse_resume_structure(doc)
    _log.debug(f"Resume sections found: {list(structure.keys())}")

    existing_text = _extract_all_text(doc).lower()
    existing_bullets = _extract_bullets(doc)
    sample_bullet_style = _detect_bullet_style(doc)

    decisions: List[EditDecision] = []

    for project in projects:
        decision = _process_project(
            project=project,
            doc=doc,
            structure=structure,
            existing_text=existing_text,
            existing_bullets=existing_bullets,
            sample_style=sample_bullet_style,
            llm=llm,
            min_confidence=min_confidence,
        )
        decisions.append(decision)
        existing_text += " " + project.primary_bullet.text.lower()

    doc.save(output_path)
    _log.success(f"Resume saved -> {output_path}")
    return decisions


def _process_project(
    project: TransformedProject,
    doc: Document,
    structure: Dict[str, int],
    existing_text: str,
    existing_bullets: List[str],
    sample_style: dict,
    llm: LLMClient,
    min_confidence: float,
) -> EditDecision:
    repo = project.repo_name

    if _is_duplicate(repo, project.primary_bullet.text, existing_text, existing_bullets):
        _log.warn(f"{repo} -> already on resume, skipping")
        return EditDecision(
            repo_name=repo,
            action="skipped_duplicate",
            reason="Project or similar bullet already exists in resume",
        )

    if project.primary_bullet.confidence < min_confidence:
        _log.warn(f"{repo} -> confidence {project.primary_bullet.confidence:.2f} below threshold")
        return EditDecision(
            repo_name=repo,
            action="skipped_low_quality",
            reason=f"Confidence {project.primary_bullet.confidence:.2f} < {min_confidence}",
        )

    bullet_text = _style_check(project, existing_bullets[:3], llm)

    section_keys = ["project", "experience", "work"]
    if project.section_target.lower() == "experience":
        section_keys = ["experience", "work", "project"]

    section_idx = _find_section(doc, section_keys)

    if section_idx is None:
        _log.warn(f"No matching section found for {repo}, appending new Projects section")
        _append_new_section(doc, "Projects", bullet_text, sample_style)
    else:
        _insert_bullet(doc, section_idx, bullet_text, sample_style)

    skills_added = _update_skills(doc, project.skills_to_add, existing_text)
    _log.success(f"{repo} -> inserted bullet, +{len(skills_added)} skills")

    return EditDecision(
        repo_name=repo,
        action="added",
        reason="Inserted as new project bullet with style check",
        bullet_used=bullet_text,
        skills_added=skills_added,
    )


def _para_style_name(para) -> str:
    """Safely get paragraph style name — returns '' if style is None."""
    if para.style is None or para.style.name is None:
        return ""
    return para.style.name.lower()


def _is_bold_header(para) -> bool:
    """True if the paragraph looks like a section header via bold formatting (no heading style)."""
    text = para.text.strip()
    if not text or len(text) > 50:
        return False
    runs = [r for r in para.runs if r.text.strip()]
    return bool(runs) and all(r.bold for r in runs)


def _parse_resume_structure(doc: Document) -> Dict[str, int]:
    structure = {}
    for i, para in enumerate(doc.paragraphs):
        style = _para_style_name(para)
        if style.startswith("heading") or _is_bold_header(para):
            structure[para.text.strip().lower()] = i
    return structure


def _find_section(doc: Document, keywords: List[str]) -> Optional[int]:
    for i, para in enumerate(doc.paragraphs):
        style = _para_style_name(para)
        is_section = "heading" in style or _is_bold_header(para)
        if is_section:
            text = para.text.lower()
            if any(kw in text for kw in keywords):
                return i
    return None


def _extract_all_text(doc: Document) -> str:
    return "\n".join(p.text for p in doc.paragraphs)


def _is_bullet_para(para) -> bool:
    """True if the paragraph is a resume bullet line (not a heading or long prose block)."""
    text = para.text.strip()
    if not text or len(text) < 20:
        return False
    style = _para_style_name(para)
    if "heading" in style:
        return False
    # List-style paragraphs
    if any(kw in style for kw in _BULLET_STYLE_KEYWORDS):
        return True
    # Paragraphs starting with common bullet characters
    if text[0] in _BULLET_CHARS:
        return True
    # Short, actionable lines typical of resume bullets (20–200 chars)
    # Exclude long prose (job summaries, skill paragraphs)
    return len(text) <= 200


def _extract_bullets(doc: Document) -> List[str]:
    bullets = []
    for para in doc.paragraphs:
        if _is_bullet_para(para):
            bullets.append(para.text.strip())
    return bullets[:20]


def _detect_bullet_style(doc: Document) -> dict:
    for para in doc.paragraphs:
        if para.text.strip() and not _para_style_name(para).startswith("heading"):
            for run in para.runs:
                size = run.font.size
                bold = run.font.bold
                sname = para.style.name if para.style else "Normal"
                return {"size": size, "bold": bold or False, "style_name": sname}
    return {"size": Pt(11), "bold": False, "style_name": "Normal"}


def _normalize_skill(skill: str) -> str:
    """Lowercase and strip all non-alphanumeric chars for comparison."""
    return re.sub(r"[^a-z0-9]", "", skill.lower())


def _is_duplicate(
    repo_name: str,
    bullet: str,
    existing_text: str,
    existing_bullets: List[str],
) -> bool:
    if repo_name.lower().replace("-", " ").replace("_", " ") in existing_text:
        return True
    bullet_words = set(re.sub(r"[^a-z0-9 ]", "", bullet.lower()).split())
    # Strip common stop words that inflate overlap
    stop = {"a", "an", "the", "and", "to", "of", "for", "in", "with", "on", "by", "at"}
    bullet_words -= stop
    if not bullet_words:
        return False
    for eb in existing_bullets:
        eb_words = set(re.sub(r"[^a-z0-9 ]", "", eb.lower()).split()) - stop
        overlap = len(bullet_words & eb_words)
        if overlap > 0 and overlap / max(len(bullet_words), 1) > 0.4:
            return True
    return False


def _style_check(
    project: TransformedProject,
    sample_bullets: List[str],
    llm: LLMClient,
) -> str:
    if not sample_bullets:
        return project.primary_bullet.text

    samples = "\n".join(f"- {b}" for b in sample_bullets[:3])
    user_prompt = f"""Existing resume bullets (for style reference):
{samples}

New bullet to check:
"{project.primary_bullet.text}"

Does it match the style? Rewrite if needed.
Return JSON: {{"bullet": "...", "changed": true/false, "reason": "..."}}"""

    from pydantic import BaseModel as PydanticBase

    class _StyleResp(PydanticBase):
        bullet: str
        changed: bool = False
        reason: str = ""

    try:
        resp = llm.structured(STYLE_CHECK_SYSTEM, user_prompt, _StyleResp)
        if resp.changed:
            _log.debug(f"Style adjusted: {resp.reason}")
        return resp.bullet
    except Exception:
        return project.primary_bullet.text


def _add_paragraph_safe(doc: Document, style_name: str):
    """Add a paragraph, falling back through styles to the document default."""
    for name in (style_name, "Normal", None):
        try:
            return doc.add_paragraph(style=name)
        except KeyError:
            continue
    return doc.add_paragraph()


def _insert_bullet(doc: Document, section_idx: int, text: str, sample_style: dict) -> None:
    insert_after_idx = section_idx
    for i, para in enumerate(doc.paragraphs[section_idx + 1:], start=section_idx + 1):
        if "heading" in _para_style_name(para) or _is_bold_header(para):
            break
        insert_after_idx = i

    ref_para = doc.paragraphs[insert_after_idx]
    ref_style = _para_style_name(ref_para)
    if ref_style.startswith("heading") or _is_bold_header(ref_para):
        style_name = "Normal"
    else:
        style_name = ref_para.style.name if ref_para.style else "Normal"

    new_para = _add_paragraph_safe(doc, style_name)
    new_para._element.getparent().remove(new_para._element)
    ref_para._element.addnext(new_para._element)

    run = new_para.add_run(text)
    if sample_style.get("size"):
        run.font.size = sample_style["size"]
    run.font.bold = sample_style.get("bold", False)


def _append_new_section(doc: Document, title: str, first_bullet: str, sample_style: dict) -> None:
    doc.add_heading(title, level=2)
    para = _add_paragraph_safe(doc, "List Bullet")
    run = para.add_run(first_bullet)
    if sample_style.get("size"):
        run.font.size = sample_style["size"]


def _update_skills(doc: Document, new_skills: List[str], existing_text: str) -> List[str]:
    skills_idx = _find_section(doc, ["skill", "technologies", "tech stack", "tools"])
    if skills_idx is None:
        return []

    skills_text = ""
    for para in doc.paragraphs[skills_idx + 1:]:
        if "heading" in _para_style_name(para) or _is_bold_header(para):
            break
        skills_text += " " + para.text

    existing_normalized = {_normalize_skill(s) for s in re.split(r"[,|•\n]", skills_text) if s.strip()}

    truly_new = []
    for s in new_skills:
        norm = _normalize_skill(s)
        if norm and norm not in existing_normalized:
            truly_new.append(s)
            existing_normalized.add(norm)  # prevent adding duplicates within the same run

    if not truly_new:
        return []

    last_para = doc.paragraphs[skills_idx + 1] if skills_idx + 1 < len(doc.paragraphs) else None
    for para in doc.paragraphs[skills_idx + 1:]:
        if "heading" in _para_style_name(para) or _is_bold_header(para):
            break
        last_para = para

    if last_para and last_para.runs:
        last_para.runs[-1].text += ", " + ", ".join(truly_new)
    elif last_para:
        last_para.add_run(", ".join(truly_new))

    return truly_new
